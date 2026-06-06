import io

from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> dict:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        raise ValueError("Unable to open file. It may be corrupted or not a valid DOCX.")

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)

    if len(full_text.strip()) < 10:
        raise ValueError("No readable text found in this document.")

    return {
        "text": full_text,
        "page_count": max(1, len(paragraphs) // 40),
    }
